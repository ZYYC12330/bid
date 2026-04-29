from __future__ import annotations

import argparse
import base64
import json
import shutil
from pathlib import Path
from typing import Any
from urllib.parse import urlparse
from urllib.request import urlopen
from copy import deepcopy

from docx import Document
from docx.shared import Emu
from docx.text.paragraph import Paragraph

try:
    from .extract_bid_template_items import extract_rect_placeholder_metadata, first_descendant, shape_attr
except ImportError:
    from extract_bid_template_items import extract_rect_placeholder_metadata, first_descendant, shape_attr


DEFAULT_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent
    / "4.15测试-箱室类"
    / "输出模版"
    / "GKZH-25ZXH856-401.88项目箱室设备采购-招标文件-发布版1128_商务标模版.docx"
)
DEFAULT_FILLED_OUTPUT_DIR = DEFAULT_TEMPLATE_PATH.parent.parent / "填充结果"
DEFAULT_ITEMS_PATH = DEFAULT_TEMPLATE_PATH.with_suffix(".待填项清单.json")
DEFAULT_ANSWERS_PATH = DEFAULT_TEMPLATE_PATH.with_suffix(".AI填充结果.json")
DEFAULT_OUTPUT_PATH = DEFAULT_FILLED_OUTPUT_DIR / f"{DEFAULT_TEMPLATE_PATH.stem}_已填充.docx"
REPO_ROOT = Path(__file__).resolve().parent.parent
IMAGE_FRONT_MARKERS = ("正面", "前面", "首页", "人像面", "front")
IMAGE_BACK_MARKERS = ("反面", "背面", "背页", "国徽", "签发机关", "有效期限", "back")


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def answer_map(ai_answers: dict[str, Any]) -> dict[str, dict[str, Any]]:
    answers = ai_answers.get("answers", [])
    return {answer["item_id"]: answer for answer in answers if answer.get("item_id")}


def replace_text_in_runs(paragraph: Paragraph, old: str, new: str) -> bool:
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True
    return False


def is_parentheses_placeholder(value: str) -> bool:
    stripped = value.strip()
    return (
        len(stripped) >= 2
        and ((stripped.startswith("（") and stripped.endswith("）")) or (stripped.startswith("(") and stripped.endswith(")")))
    )


def fill_paragraph_underlined_blank(
    paragraph: Paragraph,
    blank_run_index: int | None,
    value: str,
    *,
    start_run_index: int = 0,
) -> bool:
    if blank_run_index is not None and 0 <= blank_run_index < len(paragraph.runs):
        run = paragraph.runs[blank_run_index]
        if run.underline or not run.text.strip() or "_" in run.text:
            run.text = value
            return True

    for run in paragraph.runs[start_run_index:]:
        if run.text.strip() and "_" not in run.text:
            continue
        if run.underline or "_" in run.text:
            run.text = value
            return True
    return False


def fill_following_underlined_blank_after_text(paragraph: Paragraph, anchor_text: str, value: str) -> bool:
    anchor_seen = False
    for run_index, run in enumerate(paragraph.runs):
        if anchor_text in run.text:
            anchor_seen = True
            tail = run.text.split(anchor_text, 1)[1]
            if tail and not tail.strip() and (run.underline or "_" in tail):
                run.text = run.text.replace(tail, value, 1)
                return True
            return fill_paragraph_underlined_blank(
                paragraph,
                None,
                value,
                start_run_index=run_index + 1,
            )
    if anchor_seen:
        return fill_paragraph_underlined_blank(paragraph, None, value)
    return False


def fill_paragraph_placeholder(paragraph: Paragraph, placeholder: str, value: str) -> None:
    if is_parentheses_placeholder(placeholder):
        fill_following_underlined_blank_after_text(paragraph, placeholder, value)
        return
    if replace_text_in_runs(paragraph, placeholder, value):
        return
    paragraph.text = paragraph.text.replace(placeholder, value, 1)


def fill_paragraph_colon(paragraph: Paragraph, label_text: str, value: str) -> None:
    for separator in ("：", ":"):
        token = f"{label_text}{separator}"
        for run_index, run in enumerate(paragraph.runs):
            if token not in run.text:
                continue

            if run.text.endswith(token):
                following_runs = paragraph.runs[run_index + 1 :]
                for next_index, next_run in enumerate(following_runs):
                    if next_run.text and next_run.text.strip():
                        break
                    if next_run.underline:
                        has_following_text = any(
                            later_run.text.strip() for later_run in following_runs[next_index + 1 :]
                        )
                        next_run.text = f"{value} " if has_following_text else value
                        return

            run.text = run.text.replace(token, f"{token}{value}", 1)
            return

    if paragraph.runs:
        paragraph.runs[-1].text = f"{paragraph.runs[-1].text}{value}"
    else:
        paragraph.text = f"{paragraph.text}{value}"


def write_cell_text(cell, value: str) -> None:
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = value
        for paragraph in cell.paragraphs[1:]:
            for run in paragraph.runs:
                run.text = ""
    else:
        cell.text = value


def _image_value_to_path(value: Any, work_dir: Path) -> Path | None:
    if isinstance(value, dict):
        for key in ("path", "image_path", "imagePath", "file_path", "filePath", "url", "image_url", "imageUrl"):
            if key in value:
                return _image_value_to_path(value[key], work_dir)
        return None

    if not isinstance(value, str):
        return None

    raw_value = value.strip()
    if not raw_value:
        return None

    if raw_value.startswith("data:image/"):
        header, _, encoded = raw_value.partition(",")
        if not encoded:
            return None
        image_format = header.split(";", 1)[0].rsplit("/", 1)[-1] or "png"
        target_path = work_dir / f"image_value.{image_format}"
        target_path.parent.mkdir(parents=True, exist_ok=True)
        target_path.write_bytes(base64.b64decode(encoded))
        return target_path

    parsed = urlparse(raw_value)
    if parsed.scheme in ("http", "https") and parsed.netloc:
        suffix = Path(parsed.path).suffix or ".png"
        target_path = work_dir / f"image_value{suffix}"
        target_path.parent.mkdir(parents=True, exist_ok=True)
        with urlopen(raw_value) as response:
            content = response.read()
        if not content:
            return None
        target_path.write_bytes(content)
        return target_path

    if raw_value.startswith("/knowledge-images/"):
        path = (REPO_ROOT / "runtime" / "knowledge_images" / Path(raw_value).name).resolve()
        if path.exists():
            return path

    path = Path(raw_value).expanduser().resolve()
    if path.exists():
        return path

    frontend_asset = raw_value.removeprefix("./").lstrip("/")
    if frontend_asset.startswith("assets/"):
        path = (REPO_ROOT / "frontend" / frontend_asset).resolve()
        if path.exists():
            return path

    return None


def _image_side_from_item(item: dict[str, Any]) -> str | None:
    field_name = str(item.get("field_name") or item.get("fieldName") or item.get("name") or "")
    normalized = field_name.lower()
    if any(marker.lower() in normalized for marker in IMAGE_BACK_MARKERS):
        return "back"
    if any(marker.lower() in normalized for marker in IMAGE_FRONT_MARKERS):
        return "front"

    locator = item.get("locator") or {}
    if locator.get("block_type") != "image_placeholder":
        return None
    try:
        shape_index = int(locator.get("shape_index"))
    except (TypeError, ValueError):
        return None
    if shape_index == 0:
        return "front"
    if shape_index == 1:
        return "back"
    return None


def _image_value_for_item(item: dict[str, Any], value: Any) -> Any:
    if not isinstance(value, dict):
        return value

    side = _image_side_from_item(item)
    if side is None:
        return value

    side_keys = (
        ("front", "front_path", "frontPath", "front_url", "frontUrl", "正面", "正面图片")
        if side == "front"
        else ("back", "back_path", "backPath", "back_url", "backUrl", "反面", "反面图片")
    )
    for key in side_keys:
        if key in value:
            return value[key]

    images = value.get("images")
    if isinstance(images, dict):
        for key in side_keys:
            if key in images:
                return images[key]

    return value


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _placeholder_shapes(paragraph: Paragraph) -> list[Any]:
    return [
        shape
        for shape in paragraph._p.xpath(".//*[local-name()='anchor' or local-name()='inline']")
        if _is_empty_rect_shape(shape)
    ]


def _is_empty_rect_shape(shape: Any) -> bool:
    geometry = first_descendant(shape, "prstGeom")
    if geometry is None or geometry.get("prst") != "rect":
        return False
    text_nodes = [child.text or "" for child in shape.iter() if str(child.tag).rsplit("}", 1)[-1] == "t"]
    return not "".join(text_nodes).strip()


def _shape_doc_pr_id(shape: Any) -> str:
    doc_pr = first_descendant(shape, "docPr")
    return str(shape_attr(doc_pr, "id") or "") if doc_pr is not None else ""


def fill_image_placeholder(document: Document, locator: dict[str, Any], image_path: Path) -> bool:
    paragraph = document.paragraphs[locator["paragraph_index"]]
    shapes = _placeholder_shapes(paragraph)
    shape: Any | None = None

    doc_pr_id = str(locator.get("doc_pr_id") or "")
    if doc_pr_id:
        for candidate in shapes:
            if _shape_doc_pr_id(candidate) == doc_pr_id:
                shape = candidate
                break

    if shape is None:
        shape_index = int(locator.get("shape_index") or 0)
        if shape_index < 0 or shape_index >= len(shapes):
            return False
        shape = shapes[shape_index]

    width_emu = int(locator.get("width_emu") or (extract_rect_placeholder_metadata(shape) or {}).get("width_emu") or 0)
    height_emu = int(locator.get("height_emu") or (extract_rect_placeholder_metadata(shape) or {}).get("height_emu") or 0)
    if width_emu <= 0 or height_emu <= 0:
        return False

    scratch_paragraph = document.add_paragraph()
    scratch_run = scratch_paragraph.add_run()
    picture = scratch_run.add_picture(str(image_path), width=Emu(width_emu), height=Emu(height_emu))
    picture_graphic = first_descendant(picture._inline, "graphic")
    if picture_graphic is None:
        _remove_paragraph(scratch_paragraph)
        return False

    old_graphic = first_descendant(shape, "graphic")
    if old_graphic is None:
        _remove_paragraph(scratch_paragraph)
        return False

    parent = old_graphic.getparent()
    parent.replace(old_graphic, deepcopy(picture_graphic))
    _remove_paragraph(scratch_paragraph)
    return True


def fill_template(
    template_path: Path,
    item_payload: dict[str, Any],
    ai_answers: dict[str, Any],
    output_path: Path,
    *,
    missing_marker: str = "【待确认】",
) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template_path, output_path)
    document = Document(output_path)
    answers_by_id = answer_map(ai_answers)
    image_work_dir = output_path.parent / f"{output_path.stem}_images"

    for item in item_payload.get("items", []):
        locator = item.get("locator", {})
        answer = answers_by_id.get(item.get("item_id"), {})
        status = answer.get("status")
        raw_value = answer.get("value")

        block_type = locator.get("block_type")
        if block_type == "image_placeholder" or item.get("field_type") == "image":
            if status == "filled":
                image_path = _image_value_to_path(_image_value_for_item(item, raw_value), image_work_dir)
                if image_path is not None:
                    fill_image_placeholder(document, locator, image_path)
            continue

        value = str(raw_value or "")
        if status != "filled" or not value:
            value = missing_marker

        if block_type == "paragraph_placeholder":
            paragraph = document.paragraphs[locator["paragraph_index"]]
            fill_paragraph_placeholder(paragraph, locator["placeholder_text"], value)
        elif block_type == "paragraph_underlined_blank":
            paragraph = document.paragraphs[locator["paragraph_index"]]
            fill_paragraph_underlined_blank(paragraph, locator.get("blank_run_index"), value)
        elif block_type == "paragraph_colon":
            paragraph = document.paragraphs[locator["paragraph_index"]]
            fill_paragraph_colon(paragraph, locator["label_text"], value)
        elif block_type == "table_cell":
            table = document.tables[locator["table_index"]]
            cell = table.cell(locator["row_index"], locator["cell_index"])
            write_cell_text(cell, value)

    document.save(output_path)
    return output_path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="根据待填项清单和 AI JSON 结果回填商务标/技术标 Word。")
    parser.add_argument("template_path", nargs="?", type=Path, default=DEFAULT_TEMPLATE_PATH)
    parser.add_argument("--items", type=Path, default=DEFAULT_ITEMS_PATH, help="待填项清单 JSON")
    parser.add_argument("--answers", type=Path, default=DEFAULT_ANSWERS_PATH, help="AI 填充结果 JSON")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_PATH)
    parser.add_argument("--missing-marker", default="【待确认】")
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    output_path = fill_template(
        args.template_path,
        load_json(args.items),
        load_json(args.answers),
        args.output,
        missing_marker=args.missing_marker,
    )
    print(f"已填充文档: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
