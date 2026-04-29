from __future__ import annotations

import unittest
from tempfile import TemporaryDirectory
from pathlib import Path

from docx import Document

from scripts.extract_bid_template_items import (
    extract_rect_placeholder_metadata,
    extract_template_items,
    image_placeholder_field_name,
    infer_bid_deadline_component,
    infer_field_type,
    infer_source_preference,
    is_bid_deadline_date_field,
    make_item,
)
from scripts.prepare_qwen_fill_request import build_user_payload


class ExtractBidTemplateItemsTest(unittest.TestCase):
    def test_infer_source_preference_uses_allowed_kb_names(self) -> None:
        cases = {
            "类似项目业绩": ["case_performance_kb"],
            "资质证书编号": ["qualification_materials_kb"],
            "招标编号": ["tender_requirements_kb"],
            "投标人名称": ["company_profile_kb"],
        }

        for field_name, expected in cases.items():
            with self.subTest(field_name=field_name):
                self.assertEqual(infer_source_preference(field_name), expected)

    def test_bid_deadline_date_fields_use_tender_requirements(self) -> None:
        cases = {
            "年": "year",
            "月": "month",
            "日": "day",
            "日期": "full_date",
            "投标截止时间": "full_date",
        }

        for field_name, component in cases.items():
            with self.subTest(field_name=field_name):
                self.assertTrue(is_bid_deadline_date_field(field_name))
                self.assertEqual(infer_bid_deadline_component(field_name), component)
                self.assertEqual(infer_field_type(field_name), "date_or_period")
                self.assertEqual(infer_source_preference(field_name), ["tender_requirements_kb"])

    def test_make_item_marks_bid_deadline_fill_rule(self) -> None:
        item = make_item(
            item_id="business_001",
            template_type="business",
            section="封面",
            field_name="年",
            placeholder_text="（年）",
            prompt_hint="填写段落中的占位符：（年）",
            locator={
                "block_type": "paragraph_placeholder",
                "paragraph_index": 0,
                "placeholder_text": "（年）",
            },
        )

        self.assertEqual(item["fill_rule"]["type"], "tender_bid_deadline")
        self.assertEqual(item["fill_rule"]["component"], "year")

    def test_qwen_payload_includes_bid_deadline_date_rule(self) -> None:
        payload = build_user_payload({"items": []})

        self.assertIn("bid_deadline_date_rule", payload["writing_rules"])
        self.assertIn("投标截止时间", payload["writing_rules"]["bid_deadline_date_rule"])

    def test_extracts_xxx_and_20xx_date_placeholders(self) -> None:
        with TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "template.docx"
            document = Document()
            document.add_paragraph("投标人（盖单位公章）：XXX")
            document.add_paragraph("20XX 年 XX 月 XX 日")
            document.save(path)

            payload = extract_template_items(path, template_type="business")

        items = payload["items"]
        self.assertEqual(len(items), 2)
        self.assertEqual(items[0]["field_name"], "投标人")
        self.assertEqual(items[0]["placeholder_text"], "XXX")
        self.assertEqual(items[0]["locator"]["block_type"], "paragraph_placeholder")
        self.assertEqual(items[1]["field_name"], "日期")
        self.assertEqual(items[1]["placeholder_text"], "20XX 年 XX 月 XX 日")
        self.assertEqual(items[1]["field_type"], "date_or_period")

    def test_parentheses_text_is_hint_for_following_underlined_blank(self) -> None:
        with TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "template.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("投标人（投标人名称）")
            blank_run = paragraph.add_run("      ")
            blank_run.underline = True
            document.add_paragraph("（投标人名称）")
            document.save(path)

            payload = extract_template_items(path, template_type="business")

        items = payload["items"]
        self.assertEqual(len(items), 1)
        self.assertEqual(items[0]["field_name"], "投标人名称")
        self.assertEqual(items[0]["placeholder_text"], "（投标人名称）")
        self.assertEqual(items[0]["locator"]["block_type"], "paragraph_underlined_blank")
        self.assertEqual(items[0]["locator"]["blank_run_index"], 1)

    def test_detects_empty_rect_shape_as_image_placeholder_metadata(self) -> None:
        from docx.oxml import parse_xml

        shape = parse_xml(
            """
            <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                       xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
              <wp:extent cx="2486025" cy="1390650"/>
              <wp:docPr id="2" name="矩形 2"/>
              <a:graphic>
                <a:graphicData>
                  <wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                    <wps:spPr><a:prstGeom prst="rect"/></wps:spPr>
                    <wps:txbx><w:txbxContent><w:p/></w:txbxContent></wps:txbx>
                  </wps:wsp>
                </a:graphicData>
              </a:graphic>
            </wp:anchor>
            """
        )

        metadata = extract_rect_placeholder_metadata(shape)

        self.assertEqual(metadata["doc_pr_id"], "2")
        self.assertEqual(metadata["doc_pr_name"], "矩形 2")
        self.assertEqual(metadata["width_emu"], 2486025)
        self.assertEqual(metadata["height_emu"], 1390650)

    def test_image_placeholder_field_name_uses_id_card_sides(self) -> None:
        context = "附：法定代表人（单位负责人）身份证复印件。"

        self.assertEqual(image_placeholder_field_name(context, 1), "法定代表人（单位负责人）身份证复印件正面")
        self.assertEqual(image_placeholder_field_name(context, 2), "法定代表人（单位负责人）身份证复印件反面")


if __name__ == "__main__":
    unittest.main()
