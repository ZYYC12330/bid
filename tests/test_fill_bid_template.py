from __future__ import annotations

import base64
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.oxml import parse_xml

from scripts.fill_bid_template import (
    _image_value_for_item,
    _image_value_to_path,
    fill_paragraph_colon,
    fill_paragraph_placeholder,
    fill_paragraph_underlined_blank,
    fill_template,
)


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


class FillBidTemplateTest(unittest.TestCase):
    def test_image_value_to_path_resolves_frontend_static_assets(self) -> None:
        path = _image_value_to_path("./assets/legal-representative-id-front.jpg", Path("/tmp"))

        self.assertIsNotNone(path)
        self.assertTrue(path.exists())
        self.assertTrue(str(path).endswith("frontend/assets/legal-representative-id-front.jpg"))

    def test_image_value_to_path_resolves_uploaded_knowledge_image_url(self) -> None:
        with TemporaryDirectory() as tmpdir:
            image_dir = Path("runtime/knowledge_images")
            image_dir.mkdir(parents=True, exist_ok=True)
            image_path = image_dir / "unit-test-image.png"
            image_path.write_bytes(PNG_1X1)
            try:
                path = _image_value_to_path("/knowledge-images/unit-test-image.png", Path(tmpdir))
            finally:
                image_path.unlink(missing_ok=True)

        self.assertIsNotNone(path)
        self.assertTrue(str(path).endswith("runtime/knowledge_images/unit-test-image.png"))

    def test_image_value_for_item_selects_side_from_left_right_placeholder(self) -> None:
        shared_value = {"front": "front.png", "back": "back.png"}

        front_value = _image_value_for_item(
            {
                "field_name": "法定代表人（单位负责人）身份证复印件",
                "locator": {"block_type": "image_placeholder", "shape_index": 0},
            },
            shared_value,
        )
        back_value = _image_value_for_item(
            {
                "field_name": "法定代表人（单位负责人）身份证复印件",
                "locator": {"block_type": "image_placeholder", "shape_index": 1},
            },
            shared_value,
        )

        self.assertEqual(front_value, "front.png")
        self.assertEqual(back_value, "back.png")

    def test_fill_paragraph_colon_preserves_underlined_blank_run(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("投标人：")
        blank_run = paragraph.add_run("      ")
        blank_run.underline = True
        paragraph.add_run("（盖章）")

        fill_paragraph_colon(paragraph, "投标人", "航天晨光股份有限公司")

        self.assertEqual(paragraph.text, "投标人：航天晨光股份有限公司 （盖章）")
        self.assertEqual(paragraph.runs[1].text, "航天晨光股份有限公司 ")
        self.assertTrue(paragraph.runs[1].underline)

    def test_fill_parentheses_placeholder_without_following_blank_is_not_replaced(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run("（年）")
        run.underline = True

        fill_paragraph_placeholder(paragraph, "（年）", "2026")

        self.assertEqual(paragraph.text, "（年）")
        self.assertTrue(paragraph.runs[0].underline)

    def test_fill_parentheses_placeholder_writes_following_underlined_blank(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("投标人（投标人名称）")
        blank_run = paragraph.add_run("      ")
        blank_run.underline = True

        fill_paragraph_placeholder(paragraph, "（投标人名称）", "航天晨光股份有限公司")

        self.assertEqual(paragraph.text, "投标人（投标人名称）航天晨光股份有限公司")
        self.assertEqual(paragraph.runs[0].text, "投标人（投标人名称）")
        self.assertEqual(paragraph.runs[1].text, "航天晨光股份有限公司")
        self.assertTrue(paragraph.runs[1].underline)

    def test_fill_single_parentheses_placeholder_does_not_replace_hint_text(self) -> None:
        document = Document()
        paragraph = document.add_paragraph("（投标人名称）")

        fill_paragraph_placeholder(paragraph, "（投标人名称）", "航天晨光股份有限公司")

        self.assertEqual(paragraph.text, "（投标人名称）")

    def test_fill_paragraph_underlined_blank_uses_locator_run_index(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("邮政编码")
        blank_run = paragraph.add_run("      ")
        blank_run.underline = True

        fill_paragraph_underlined_blank(paragraph, 1, "211100")

        self.assertEqual(paragraph.text, "邮政编码211100")
        self.assertTrue(paragraph.runs[1].underline)

    def test_fill_template_replaces_image_placeholder_shape_with_picture(self) -> None:
        with TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            template_path = tmp_path / "template.docx"
            image_path = tmp_path / "front.png"
            output_path = tmp_path / "filled.docx"
            image_path.write_bytes(PNG_1X1)

            document = Document()
            paragraph = document.add_paragraph("法定代表人（单位负责人）身份证复印件正面")
            paragraph._p.append(
                parse_xml(
                    """
                    <w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                         xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                         xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                         xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                      <w:drawing>
                        <wp:anchor>
                          <wp:extent cx="2000000" cy="1000000"/>
                          <wp:docPr id="9" name="矩形 9"/>
                          <a:graphic>
                            <a:graphicData>
                              <wps:wsp>
                                <wps:spPr><a:prstGeom prst="rect"/></wps:spPr>
                                <wps:txbx><w:txbxContent><w:p/></w:txbxContent></wps:txbx>
                              </wps:wsp>
                            </a:graphicData>
                          </a:graphic>
                        </wp:anchor>
                      </w:drawing>
                    </w:r>
                    """
                )
            )
            document.save(template_path)

            fill_template(
                template_path,
                {
                    "items": [
                        {
                            "item_id": "business_001",
                            "field_type": "image",
                            "locator": {
                                "block_type": "image_placeholder",
                                "paragraph_index": 0,
                                "shape_index": 0,
                                "doc_pr_id": "9",
                                "width_emu": 2000000,
                                "height_emu": 1000000,
                            },
                        }
                    ]
                },
                {"answers": [{"item_id": "business_001", "status": "filled", "value": str(image_path)}]},
                output_path,
            )

            with zipfile.ZipFile(output_path) as docx_zip:
                document_xml = docx_zip.read("word/document.xml").decode("utf-8")
                media_files = [name for name in docx_zip.namelist() if name.startswith("word/media/")]

            self.assertIn("<pic:pic", document_xml)
            self.assertNotIn("<wps:wsp", document_xml)
            self.assertTrue(media_files)


if __name__ == "__main__":
    unittest.main()
